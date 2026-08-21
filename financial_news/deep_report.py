#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Deep report generator - triggered right after the daily report is written.

Workflow:
  1. Read today's daily report markdown (knowledge-base/金融/每日報告/YYYY-MM-DD.md)
  2. Send it to Gemini for a deeper multi-asset analysis
  3. Render the analysis into a Word .docx + Markdown (knowledge-base/金融/深度報告/)
  4. Publish to GitHub Pages (via publisher/build.py); no email/LINE notifications

Usage:
  python deep_report.py              # deep report for today
  python deep_report.py --day 2026-08-04
  python deep_report.py --skip-email --skip-line   # kept for backward compatibility
"""
import os
import re
import sys
import argparse
import logging
from urllib.parse import quote_plus
from datetime import datetime, timedelta, timezone

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
except ImportError:
    Document = None

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import log_util

REPORT_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "knowledge-base", "金融", "每日報告",
)
DEEP_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "..", "knowledge-base", "金融", "深度報告",
)
MODEL = os.environ.get("GEMINI_MODEL", "gemini-3.5-flash-lite")

LOG_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs", "deep_report.log")
logger = log_util.get_logger(__name__, LOG_FILE)

SYSTEM_PROMPT = """你是一名資深的國際金融分析師。使用者會給你一份當日的《金融重點報告》，也可能附加近兩日的報告（供跨日相關性分析）與一份分析師人工撰寫的報告（Big Pickle，特色是涵蓋台股夜盤變動、跨日因果與資金輪動視角）。

請產出一份【深度分析報告】，以繁體中文、Markdown 格式輸出。融合兩者特色：
- 保留 Gemini 的結構化分類（市場總覽／分類展望／風險機會）
- 融入 Big Pickle 的即時感：台股夜盤變動、跨日相關性（近三日因果）、資金輪動、具體投資建議

結構如下：

## 一、市場總覽
針對股/債/匯/商品四大市場，給出今日整體方向與一句話總結。

## 二、行情快照與台股夜盤解析
系統會額外提供一份【今日最新行情快照（唯一可信來源）】：本節必須以該快照的數據與「快照資料時間」為準，嚴禁引用近兩日報告或分析師人工報告中的行情價格或快照時間作為當前行情。並重點解析台股夜盤／台指期的即時變動，以及與美股的背離或連動。

## 三、重大事件深度解讀
針對報告中的重大事件（severity 較高者），逐一拆解：
- 事件核心
- 對股/債/匯/商品的傳導路徑
- 關聯標的（台股代號/美股代號）
- 未來一週可能的走勢

## 四、跨日相關性分析
如有近兩日報告，分析近三日市場演變的因果閉環與主軸（例如地緣、債市、資金輪動、台股背離），點出趨勢如何延續或轉折。

## 五、分類展望
分別就股市、債市、匯市、商品(油金)、地緣政治，給出短(1週)/中(1月)/長(1季)展望。資訊不足的項目要誠實說明，不要編造。

## 六、風險與機會
列出本報告中最值得關注的風險（警示）與機會（看好），各列 3-5 項，每項含標的與理由。

## 七、操作建議摘要
以投資者視角，給出 3-5 條具體、可執行的觀察重點（不是投資建議），並註明需要盯盤的關鍵指標（如油價、美債殖利率、台指期夜盤）。

## 八、免責聲明
本報告由 AI 自動生成，僅供參考，不構成投資建議。

格式要求（重要）：
- 每個「##」大標題下，先列出**次重點標題**（用 `###` 或 `- **次重點標題**`），其下再給該次重點的**細部說明**一到數行。
- 全文的重要結論、關鍵事項、警示重點請用 **粗體（Markdown 雙星號）** 標示，方便快速掃讀。
- 開頭不要重複「報告時間」（系統會自動加）。

要求：內容具體、有數據依據（引用報告中的數字與事件）、避免空泛形容詞。全篇約 1000-1500 字。直接輸出 Markdown，不要用程式碼圍欄包裹。"""


def _read_key():
    api_key = os.environ.get("GEMINI_API_KEY") or ""
    if not api_key:
        kf = os.path.join(os.path.expanduser("~"), ".gemini_api_key")
        if os.path.exists(kf):
            api_key = open(kf, encoding="utf-8").read().strip()
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY not found")
    return api_key


def read_daily_report(day):
    path = os.path.join(REPORT_DIR, f"{day}.md")
    if not os.path.exists(path):
        return None, path
    with open(path, encoding="utf-8") as f:
        return f.read(), path


def _extract_doc_text(path):
    """Extract readable text from a legacy Word 97-2003 .doc (OLE2).

    The body is stored as UTF-16LE in the compound file; we decode the whole
    blob and slice out the readable body, dropping the binary style tables.
    """
    with open(path, "rb") as f:
        data = f.read()
    text = data.decode("utf-16-le", errors="ignore")
    m = re.search(r"(綜合深度分析報告|深度分析報告)[^\n]{0,40}", text)
    start = m.start() if m else 0
    body = text[start:]
    for marker in ("伀倀儀", "Root Entry", "標題 1 字元", "Normal.dot"):
        idx = body.find(marker)
        if idx > 200:
            body = body[:idx]
            break
    return body.strip()


def find_bigpickle_report(day):
    """Locate Big Pickle's hand-written deep report .doc for the day, if any."""
    candidates = [
        os.path.join(REPORT_DIR, f"綜合深度分析報告 {day}.doc"),
        os.path.join(REPORT_DIR, f"深度分析報告 {day}.doc"),
        os.path.join(REPORT_DIR, f"{day}.doc"),
    ]
    for c in candidates:
        if os.path.exists(c):
            return c
    return None


def _fresh_snapshot_text():
    """Pull the latest market snapshot straight from the DB (same source the
    daily report uses) so the model never relies on stale numbers found in
    previous-day reports (2026-08-20 PC3 incident: deep report quoted the
    prior day's close because prev-days text contained it)."""
    try:
        import market_data
        table = market_data.latest_table()
        if not table:
            return None
        cap_disp, age_h, stale = market_data.latest_meta()
        warn = "（注意：快照已超過 24 小時未更新，請在報告中明確標示時效限制）" if stale else ""
        return (
            "【今日最新行情快照（系統直接自資料庫取得，唯一可信來源）】\n"
            f"快照資料時間：{cap_disp}{warn}\n"
            f"{table}\n"
            "『二、行情快照』一節必須使用上表數據與時間；"
            "嚴禁引用近兩日報告或 Big Pickle 報告中的行情價格或快照時間作為當前行情。"
        )
    except Exception as e:
        print(f"[deep_report] fresh snapshot unavailable: {e}")
        return None


def _verification_pass(analysis_text):
    """Run a fact-check pass on the generated deep report. Returns (text, note)."""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=_read_key())
    check_prompt = (
        "你是一名嚴謹的金融報告審稿人。以下是一份 AI 生成的《深度分析報告》。\n"
        "請只檢查報告內容的『內部邏輯一致性』，判斷是否有：\n"
        "1. 同一數字在文中前後矛盾（例如同一指數出現兩個不同價位/漲跌幅）\n"
        "2. 明顯的日期/時間錯誤或不一致\n"
        "3. 標的（代號/名稱）在文內自相矛盾或與報告內其他描述不符\n\n"
        "【重要限制】你的知識庫可能過期，請勿使用任何外部知識判斷下列事項的真假，"
        "包括但不限於：公司是否上市、股價高低是否合理、指數點位與匯率利率是否合理、"
        "某產業趨勢是否存在。報告內容描述的是最新即時新聞，可能超出你的知識範圍。\n"
        "若某個描述『僅與外部事實可能不符』但報告內部前後一致，請不要列出為疑點。\n"
        "【判斷範例】報告寫『SpaceX 在那斯達克掛牌，代號 SPCX』。即使你以為 SpaceX "
        "是未上市私人公司，也請不要列為疑點——你的知識可能過期，這是外部事實判斷。\n"
        "只有當報告內部同時出現互相矛盾的說法（例如一處說 SpaceX 未上市、另一處說已上市）時，"
        "才可列為疑點。\n\n"
        "若發現內部矛盾，請輸出『修正清單』：每條為『【疑點】<出處> - <問題說明> - <建議修正>』。\n"
        "若無問題，請輸出『未發現矛盾』。\n"
        "請務必只基於報告內容本身的邏輯一致性判斷。\n\n"
        f"報告內容：\n{analysis_text[:12000]}"
    )
    try:
        resp = client.models.generate_content(
            model=MODEL,
            contents=check_prompt,
            config=types.GenerateContentConfig(temperature=0.0),
        )
        note = (resp.text or "").strip()
    except Exception as e:
        return analysis_text, f"核查失敗（不影響報告）: {str(e)[:120]}"
    if not note or "未發現矛盾" in note:
        return analysis_text, "核查通過：未發現明顯矛盾"
    # 有疑點 → 附註在報告開頭（不覆寫內容，供人工檢視）
    if len(note) > 2000:
        note = note[:2000] + "..."
    header = ("> ⚠️ **AI 審稿疑點（自動核查發現，請人工確認）**\n>\n"
              + "\n>\n".join("> " + ln if ln.strip() else "" for ln in note.splitlines())
              + "\n\n---\n\n")
    return header + analysis_text, f"核查發現疑點（已附註報告開頭）"


def deep_analyze(md_text, bigpickle_text=None, prev_days_text=None, major_events_text=None,
                 snapshot_text=None):
    """Call Gemini for the deep analysis. Returns markdown text."""
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=_read_key())

    parts = []
    parts.append("【今日金融重點報告】\n" + md_text[:12000])
    if snapshot_text:
        # placed immediately after today's report so it outranks the
        # prev-days text that follows
        parts.append(snapshot_text)
    if prev_days_text:
        parts.append("【近兩日金融重點報告（供跨日相關性分析）】\n" + prev_days_text)
    if bigpickle_text:
        parts.append(
            "【分析師人工報告（Big Pickle，含台股夜盤與跨日觀點，請融合其重點與視角）】\n"
            + bigpickle_text[:10000]
        )
    if major_events_text:
        parts.append(
            "【今日重大事件清單（含來源與可引用之完整網址）】\n"
            "請在報告『重大事件深度解讀』一節中，為每件重大事件附上來源說明，"
            "格式為『[來源: 媒體名稱](網址)』，務必包含完整網址供讀者點擊核對。\n"
            "若某事件標示『⚠️AI分歧』，表示兩個 AI 模型對其重要性判斷不一致，"
            "請在報告中如實標註，並以更審慎的措辭描述。\n"
            + major_events_text[:6000]
        )
    contents = "\n\n---\n\n".join(parts)

    resp = client.models.generate_content(
        model=MODEL,
        contents=contents,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=0.4,
        ),
    )
    text = (resp.text or "").strip()
    text = re.sub(r"^```(?:markdown)?\s*\n?", "", text, flags=re.M)
    text = re.sub(r"\n?```\s*$", "", text, flags=re.M)
    # Fact-check pass (free tier, does not overwrite content)
    if not os.environ.get("NO_VERIFY_PASS"):
        text, note = _verification_pass(text)
        print(f"[deep_report] 自動核查: {note}")
    return text


FONT_NAME = "華康仿宋體W4"
FONT_SIZE = Pt(12)
TITLE_COLOR = RGBColor(0x00, 0x00, 0x00)      # 一級標題黑色
SUB_TITLE_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # 次標題藍色（加深）
KEY_COLOR = RGBColor(0x00, 0x32, 0x6E)         # 重點/關鍵事項深藍色


def _set_ea_font(run, name=FONT_NAME, size=FONT_SIZE, color=None, bold=None):
    """Set the font (Chinese via eastAsia, Latin via ascii) on a run."""
    run.font.name = "Times New Roman"
    run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), name)


def _fmt_para(p, before=0, after=2):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15


def _style_heading(p, level):
    """Apply font + color to every run of a heading paragraph."""
    color = TITLE_COLOR if level <= 2 else SUB_TITLE_COLOR
    for r in p.runs:
        _set_ea_font(r, size=FONT_SIZE, color=color, bold=True)


def _add_markdown(doc, md_text):
    """Render a small markdown subset (headings, bullets, numbered, bold) into docx.

    Uses the FangSong 12pt body font with compact paragraph spacing,
    light-blue titles and dark-blue key highlights.
    """
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("#### "):
            p = doc.add_heading(line[5:], level=4)
            _style_heading(p, 4)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            _style_heading(p, 3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            _style_heading(p, 2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            _style_heading(p, 1)
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich(p, line[2:])
        elif re.match(r"^\d+[.)] ", line):
            p = doc.add_paragraph(style="List Number")
            _add_rich(p, re.sub(r"^\d+[.)] ", "", line))
        elif line.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            _set_ea_font(r)
        elif line.strip() == "---":
            continue
        else:
            _add_rich(doc.add_paragraph(), line)
        if not line.startswith(">"):
            _fmt_para(p, before=6 if line.startswith("#") else 0)


def _add_rich(paragraph, text):
    """Add text with **bold** (key highlights, dark blue) and `inline code` support."""
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _set_ea_font(paragraph.add_run(part[2:-2]), color=KEY_COLOR, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            _set_ea_font(r)
        else:
            _set_ea_font(paragraph.add_run(part))


def _docx_to_doc(docx_path, doc_path):
    """Convert a .docx to legacy .doc via Word COM."""
    try:
        import win32com.client
    except ImportError:
        return False
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        d = word.Documents.Open(docx_path, False, False)
        d.SaveAs(doc_path, FileFormat=0)  # 0 = wdFormatDocument97 (.doc)
        d.Close(False)
        return os.path.exists(doc_path)
    except Exception as e:
        print(f"[deep_report] .doc 轉換失敗: {e}")
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def write_doc(analysis_md, day, now_str, label=""):
    os.makedirs(DEEP_DIR, exist_ok=True)
    # Also persist the markdown into the KB so Obsidian can read the deep
    # report (same folder, next to the .docx/.doc files).
    md_path = os.path.join(DEEP_DIR, f"深度分析報告 {day} {label}.md").strip()
    try:
        header = f"> 報告時間：{now_str}\n\n---\n\n"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(header + analysis_md)
        print(f"[deep_report] Markdown 已存知識庫: {md_path}")
    except Exception as e:
        print(f"[deep_report] Markdown 存檔失敗: {e}")
    docx_path = os.path.join(DEEP_DIR, f"{day}.docx")
    doc = Document()
    for style in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            st = doc.styles[style]
            st.font.name = "Times New Roman"
            st.font.size = FONT_SIZE
            st.element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            st.element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        except Exception:
            pass
    # Report timestamp as the opening line (bold, key highlight)
    p = doc.add_paragraph()
    r = p.add_run(f"報告時間：{now_str}")
    _set_ea_font(r, color=TITLE_COLOR, bold=True)
    _fmt_para(p, after=4)
    _add_markdown(doc, analysis_md)
    doc.save(docx_path)

    # Convert to .doc (Word 97-2003) for final delivery
    doc_path = os.path.join(DEEP_DIR, f"深度分析報告 {day} {label}.doc").strip()
    if _docx_to_doc(docx_path, doc_path):
        return doc_path
    print("[deep_report] 改用 .docx 交付")
    return docx_path


def _prev_days_text(day):
    """Collect the previous two days' daily reports for cross-day analysis."""
    parts = []
    try:
        base = datetime.strptime(day, "%Y-%m-%d")
    except ValueError:
        return ""
    for i in (1, 2):
        prev = (base - timedelta(days=i)).strftime("%Y-%m-%d")
        md, _ = read_daily_report(prev)
        if md:
            parts.append(f"===== {prev} =====\n" + md[:8000])
    return "\n\n".join(parts)


def _major_events_with_sources(day, limit=15):
    """Pull today's important events with source links and verification status."""
    try:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        import db
        conn = db.connect()
        rows = conn.execute(
            """SELECT id, title, link, source, published, severity, sentiment,
                      impact_notes, verification, verifier_note
               FROM events
               WHERE is_noise=0 AND is_duplicate=0 AND severity>=2
                 AND substr(published,1,10)=?
               ORDER BY severity DESC, published DESC LIMIT ?""",
            (day, limit),
        ).fetchall()
        conn.close()
        lines = []
        for r in rows:
            flag = ""
            if r["verification"] == "conflict":
                flag = "  ⚠️AI分歧"
            elif r["verification"] == "verified":
                flag = "  ✅已複核"
            link = r["link"] or ""
            if "news.google.com" in link:
                link = "https://news.google.com/search?q=" + quote_plus(r["title"])
            lines.append(
                f"- [{r['severity']}] {r['title']}{flag}\n"
                f"  來源: {r['source']}\n"
                f"  網址: {link}"
            )
        return "\n".join(lines) if lines else "（今日無重大事件）"
    except Exception as e:
        return f"（重大事件來源擷取失敗: {e}）"


SLOT_LABELS = {"morning": "早上", "evening": "傍晚"}


def _slot_label(slot):
    return SLOT_LABELS.get(slot, "")


def run(day=None, slot=None, skip_email=False, skip_line=False, force=False):
    """Generate the deep report (Word + Markdown into KB; no email/LINE since 2026-08).

    skip_email/skip_line are accepted for backward compatibility but are no-ops now.
    """
    tz = timezone(timedelta(hours=8))
    now = datetime.now(tz)
    day = day or now.strftime("%Y-%m-%d")
    if slot is None:
        slot = "morning" if now.hour < 13 else "evening"
    label = _slot_label(slot)

    md_text, report_path = read_daily_report(day)
    if md_text is None:
        print(f"[deep_report] 找不到當日報告: {report_path}")
        return None

    doc_path = os.path.join(DEEP_DIR, f"深度分析報告 {day} {label}.doc")
    if os.path.exists(doc_path) and not force:
        print(f"[deep_report] {day} {label}深度報告已存在，跳過（--force 可覆蓋）: {doc_path}")
        return doc_path

    print(f"[deep_report] 讀取每日報告: {report_path} ({len(md_text)} 字)")

    bp_path = find_bigpickle_report(day)
    bigpickle_text = None
    if bp_path:
        try:
            bigpickle_text = _extract_doc_text(bp_path)
            print(f"[deep_report] 融合 Big Pickle 人工報告: {bp_path} ({len(bigpickle_text)} 字)")
        except Exception as e:
            print(f"[deep_report] Big Pickle 報告讀取失敗: {e}")

    prev_days = _prev_days_text(day)
    if prev_days:
        print(f"[deep_report] 附加近兩日報告作跨日分析 ({len(prev_days)} 字)")

    major_events = _major_events_with_sources(day)
    print(f"[deep_report] 重大事件引註資料已載入")

    snapshot_text = _fresh_snapshot_text()
    if snapshot_text:
        print("[deep_report] 已注入今日最新行情快照（唯一可信來源）")

    analysis_md = deep_analyze(md_text, bigpickle_text=bigpickle_text,
                               prev_days_text=prev_days, major_events_text=major_events,
                               snapshot_text=snapshot_text)
    now_str = now.strftime("%Y-%m-%d %H:%M（台灣時間）")
    doc_path = write_doc(analysis_md, day, now_str, label=label)
    print(f"[deep_report] {label}深度報告已產生: {doc_path}")
    print("[deep_report] 已停用 Email/LINE 寄送，報告由 GitHub Pages 發佈，Word 檔留存知識庫備份")

    logger.info("deep report done: %s", doc_path)
    return doc_path


def main():
    parser = argparse.ArgumentParser(description="Financial deep report generator")
    parser.add_argument("--day", default=None, help="YYYY-MM-DD (default today)")
    parser.add_argument("--slot", choices=["morning", "evening"], default=None,
                        help="morning/evening (default: auto by hour)")
    parser.add_argument("--force", action="store_true", help="overwrite existing doc")
    parser.add_argument("--skip-email", action="store_true")
    parser.add_argument("--skip-line", action="store_true")
    args = parser.parse_args()
    run(args.day, slot=args.slot, skip_email=args.skip_email, skip_line=args.skip_line, force=args.force)


if __name__ == "__main__":
    main()
